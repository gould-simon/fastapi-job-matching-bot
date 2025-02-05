import os
import logging
from openai import OpenAI, AsyncOpenAI
import pdfplumber
from docx import Document
from typing import List, Tuple
import asyncio
from functools import partial
from pathlib import Path

# Setup logging
logger = logging.getLogger(__name__)

# Initialize OpenAI client
openai = AsyncOpenAI(
    api_key=os.getenv("OPENAI_API_KEY")
)

async def extract_cv_text(file_path: str) -> str:
    """Extract text from CV file.
    
    Args:
        file_path: Path to CV file
        
    Returns:
        Extracted text from CV
        
    Raises:
        FileNotFoundError: If file doesn't exist
        ValueError: If file is empty or invalid
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
        
    # Check if file is empty
    if os.path.getsize(file_path) == 0:
        raise ValueError("Empty or invalid file")
        
    # Get file extension
    file_ext = Path(file_path).suffix.lower()
    
    # Extract text based on file type
    if file_ext == '.pdf':
        try:
            with pdfplumber.open(file_path) as pdf:
                if len(pdf.pages) == 0:
                    raise ValueError("Empty or invalid file")
                text = '\n'.join(page.extract_text() for page in pdf.pages)
        except Exception as e:
            logger.error(f"Error extracting text from CV: {str(e)}")
            raise
            
    elif file_ext == '.docx':
        try:
            doc = Document(file_path)
            if len(doc.paragraphs) == 0:
                raise ValueError("Empty or invalid file")
            text = '\n'.join(p.text for p in doc.paragraphs)
        except Exception as e:
            logger.error(f"Error extracting text from CV: {str(e)}")
            raise
            
    else:
        raise ValueError("Unsupported file format. Please upload a PDF or DOCX file.")
        
    if not text.strip():
        raise ValueError("Empty or invalid file")
        
    return text

async def process_cv(file_path: str) -> Tuple[str, List[float]]:
    """Process CV file and return extracted text and embedding."""
    try:
        # Check file size
        file_size = os.path.getsize(file_path)
        if file_size > 5 * 1024 * 1024:  # 5MB limit
            raise ValueError("File too large. Maximum file size is 5MB.")
            
        # Extract text from CV
        text = await extract_cv_text(file_path)
        
        # Generate embedding
        cv_embedding = await generate_cv_embedding(text)
        
        return text, cv_embedding
        
    except Exception as e:
        logger.error(f"Error processing CV: {str(e)}")
        raise

async def generate_cv_embedding(text: str) -> List[float]:
    """Generate embedding vector for CV text.
    
    Args:
        text: Text content to generate embedding for
        
    Returns:
        List of floats representing the embedding vector
        
    Raises:
        ValueError: If text is empty or invalid
    """
    if not text or not text.strip():
        raise ValueError("Cannot generate embedding for empty text")
        
    try:
        # Generate embedding using OpenAI API
        response = await openai.embeddings.create(
            model="text-embedding-ada-002",
            input=text
        )
        
        # Extract the embedding values
        embedding = response.data[0].embedding
        
        # Verify the embedding format
        if not isinstance(embedding, list) or not all(isinstance(x, float) for x in embedding):
            raise ValueError("Invalid embedding format from API")
            
        return embedding
        
    except Exception as e:
        logger.error(f"Error generating embedding: {str(e)}")
        raise 