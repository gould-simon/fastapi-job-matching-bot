import os
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document

def create_test_pdf():
    """Create a test PDF file with sample CV content."""
    pdf_path = os.path.join(os.path.dirname(__file__), 'test_files', 'test_cv.pdf')
    
    # Create directory if it doesn't exist
    os.makedirs(os.path.dirname(pdf_path), exist_ok=True)
    
    # Create PDF
    c = canvas.Canvas(pdf_path, pagesize=letter)
    c.setFont("Helvetica", 12)
    
    # Write content
    y_position = 750
    for line in [
        "John Doe",
        "Senior Audit Manager",
        "10+ years of experience",
        "EXPERIENCE",
        "Senior Audit Manager at Big4 Firm",
        "2015-Present",
        "• Led audit teams of 10+ professionals",
        "• Managed complex audit engagements",
        "EDUCATION",
        "MBA in Finance",
        "CPA License"
    ]:
        c.drawString(100, y_position, line)
        y_position -= 20
    c.save()

def create_test_docx():
    """Create a test Word document with sample CV content."""
    docx_path = os.path.join(os.path.dirname(__file__), 'test_files', 'test_cv.docx')
    
    # Create directory if it doesn't exist
    os.makedirs(os.path.dirname(docx_path), exist_ok=True)
    
    # Create Word document
    doc = Document()
    doc.add_heading('John Doe', 0)
    doc.add_paragraph('Senior Audit Manager')
    doc.add_paragraph('10+ years of experience')
    doc.add_heading('EXPERIENCE', level=1)
    doc.add_paragraph('Senior Audit Manager at Big4 Firm')
    doc.add_paragraph('2015-Present')
    p = doc.add_paragraph()
    p.add_run('• Led audit teams of 10+ professionals\n')
    p.add_run('• Managed complex audit engagements')
    doc.add_heading('EDUCATION', level=1)
    doc.add_paragraph('MBA in Finance')
    doc.add_paragraph('CPA License')
    doc.save(docx_path)

def create_test_files():
    """Create all test files needed for testing."""
    create_test_pdf()
    create_test_docx()

if __name__ == '__main__':
    create_test_files() 